from docx import Document

from ilo_health_report.report import write_report


def sample_data() -> dict:
    return {
        "generated_at": "2026-07-17T12:00:00+00:00",
        "target": "https://ilo.example.com",
        "server_status": {
            "Name": "server-01",
            "Model": "ProLiant DL380 Gen10",
            "Manufacturer": "HPE",
            "Serial number": "REDACTED",
            "Power state": "On",
            "BIOS version": "U30 v2.90",
            "Health": "OK",
        },
        "temperatures": [{"Name": "Ambient", "Reading (°C)": 22, "Upper critical (°C)": 42, "Health": "OK", "State": "Enabled"}],
        "fans": [{"Name": "Fan 1", "Reading": 18, "Units": "Percent", "Health": "OK", "State": "Enabled"}],
        "power_supplies": [{"Name": "Power Supply 1", "Model": "800W", "Serial number": "REDACTED", "Capacity (W)": 800, "Health": "OK", "State": "Enabled"}],
        "storage": [{"Name": "Smart Array", "Description": "Controller", "Health": "OK", "State": "Enabled"}],
        "memory": [{"Name": "PROC 1 DIMM 1", "Capacity (MiB)": 32768, "Type": "DDR4", "Speed (MHz)": 2933, "Health": "OK", "State": "Enabled"}],
        "processors": [{"Name": "CPU 1", "Model": "Intel Xeon", "Cores": 16, "Threads": 32, "Health": "OK", "State": "Enabled"}],
        "firmware": [{"Name": "iLO 5", "Version": "3.10", "Updateable": True, "Health": "OK", "State": "Enabled"}],
        "event_logs": [{"Log": "Integrated Management Log", "Created": "2026-07-16T10:00:00Z", "Severity": "Warning", "Message": "Example maintenance event", "Repaired": False}],
        "collection_notes": [],
    }


def test_write_report_contains_all_sections(tmp_path) -> None:
    output = write_report(sample_data(), tmp_path / "report.docx")
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for heading in (
        "Executive health summary",
        "Temperatures",
        "Fans",
        "Power supplies",
        "Storage",
        "Memory",
        "Processors",
        "Firmware",
        "Event logs",
    ):
        assert heading in text
    assert len(document.tables) == 9

