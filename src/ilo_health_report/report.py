"""Create the Microsoft Word health report."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
RED = "9B1C1C"
AMBER = "7A5A00"
GREEN = "2E6B3A"


def _set_cell_fill(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell: Any, value: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value if edge in ("top", "bottom") else 120))
        node.set(qn("w:type"), "dxa")


def _font(run: Any, *, size: float = 10.5, bold: bool = False, color: str = "222222") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _status_color(value: object) -> str:
    text = str(value).lower()
    if any(word in text for word in ("critical", "failed", "fatal")):
        return RED
    if any(word in text for word in ("warning", "degraded", "caution")):
        return AMBER
    if text in ("ok", "enabled", "standbyoffline"):
        return GREEN
    return "555555"


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after in (
        ("Title", 25, 0, 6),
        ("Heading 1", 16, 16, 7),
        ("Heading 2", 12.5, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Title"
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("HP iLO 5 Health Check Report")
    _font(run, size=8.5, color="666666")
    run = paragraph.add_run("  |  Page ")
    _font(run, size=8.5, color="666666")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _column_widths(records: list[dict[str, Any]], columns: list[str]) -> list[int]:
    """Allocate the compact-reference preset's 9360 DXA table width."""
    scores = []
    for column in columns:
        longest = max([len(column), *(len(str(row.get(column, ""))) for row in records)])
        scores.append(max(10, min(longest, 36)))
    total = sum(scores)
    widths = [max(900, round(9360 * score / total)) for score in scores]
    difference = 9360 - sum(widths)
    widths[widths.index(max(widths))] += difference
    return widths


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = width = Inches(widths[index] / 1440)
            cell._tc.get_or_add_tcPr().get_or_add_tcW().set(qn("w:w"), str(widths[index]))


def _add_table(document: Document, records: list[dict[str, Any]], empty: str) -> None:
    if not records:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(empty)
        run.italic = True
        _font(run, size=9.5, color="666666")
        return
    columns = list(records[0].keys())
    table = document.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, label in enumerate(columns):
        cell = table.rows[0].cells[index]
        _set_cell_fill(cell, BLUE)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(label))
        _font(run, size=9, bold=True, color=WHITE)
    for row_number, record in enumerate(records, start=1):
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cell = cells[index]
            _set_cell_margins(cell)
            if row_number % 2 == 0:
                _set_cell_fill(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            value = record.get(column, "")
            run = paragraph.add_run(str(value))
            color = _status_color(value) if column in ("Health", "Severity", "State") else "222222"
            _font(run, size=8.5, bold=column in ("Health", "Severity"), color=color)
    _set_table_geometry(table, _column_widths(records, columns))


def _add_summary_table(document: Document, summary: dict[str, Any]) -> None:
    records = [{"Item": key, "Value": value} for key, value in summary.items()]
    _add_table(document, records, "Server status was not available.")


def write_report(data: dict[str, Any], destination: str | Path) -> Path:
    path = Path(destination)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _style_document(document)
    _add_footer(document)

    title = document.add_paragraph(style="Title")
    title.add_run("HP iLO 5 Health Check Report")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(15)
    run = subtitle.add_run(f"Target: {data.get('target', 'Unknown')}  |  Generated: {data.get('generated_at', 'Unknown')}")
    _font(run, size=10, color="555555")

    document.add_heading("Executive health summary", level=1)
    summary = data.get("server_status", {})
    _add_summary_table(document, summary)

    sections = (
        ("Temperatures", "temperatures"),
        ("Fans", "fans"),
        ("Power supplies", "power_supplies"),
        ("Storage", "storage"),
        ("Memory", "memory"),
        ("Processors", "processors"),
        ("Firmware", "firmware"),
        ("Event logs", "event_logs"),
    )
    for heading, key in sections:
        document.add_heading(heading, level=1)
        _add_table(document, data.get(key, []), f"No {heading.lower()} data was returned.")

    notes = data.get("collection_notes", [])
    if notes:
        document.add_heading("Collection notes", level=1)
        for note in notes:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(str(note))

    document.save(path)
    return path
